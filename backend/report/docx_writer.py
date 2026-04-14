"""Generate Amazon selection DOCX report — professional styling."""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Brand palette ──
BRAND = "1B3A5C"        # deep navy
BRAND_LIGHT = "2B6CB0"  # medium blue
ACCENT = "E67E22"       # warm orange
HEADER_BG = "1B3A5C"    # navy header
HEADER_FG = "FFFFFF"     # white header text
ROW_EVEN = "F8FAFC"     # very light blue-gray
ROW_ODD = "FFFFFF"       # white
HIGHLIGHT_BG = "FFF8F0"  # warm highlight
BORDER = "D1D5DB"
MUTED = "6B7280"
SUCCESS = "059669"
WARNING = "D97706"
DANGER = "DC2626"


# ── XML helpers ──
def _set_cell_bg(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_cell_margins(cell, top=50, left=80, bottom=50, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_table_borders(table, color=BORDER, sz="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        borders.append(b)
    tblPr.append(borders)


def _set_run(run, *, size=10, bold=False, color=None, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_horizontal_line(paragraph, color=BRAND, sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── High-level builders ──
def _add_table(doc, headers, rows, first_col_left=True):
    """Add a styled table. first_col_left=True aligns first column left, rest center."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(t)

    # Header row
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(str(h))
        _set_run(run, size=8.5, bold=True, color=HEADER_FG)
        _set_cell_bg(hdr[i], HEADER_BG)
        _set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body rows
    for ri, r in enumerate(rows):
        rc = t.add_row().cells
        bg = ROW_EVEN if ri % 2 == 0 else ROW_ODD
        for i, v in enumerate(r):
            rc[i].text = ""
            p = rc[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (first_col_left and i == 0) else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            text = "—" if v is None or str(v).strip() in ("", "None") else str(v)
            run = p.add_run(text)
            _set_run(run, size=8.5)
            _set_cell_bg(rc[i], bg)
            _set_cell_margins(rc[i])
            rc[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ── Pagination control ──
    total_rows = len(t.rows)
    for idx, row in enumerate(t.rows):
        trPr = row._tr.get_or_add_trPr()
        # Prevent rows from splitting across pages
        cantSplit = OxmlElement("w:cantSplit")
        trPr.append(cantSplit)

    # Repeat header row on each page for large tables
    hdr_trPr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    hdr_trPr.append(tblHeader)

    # Small tables (≤5 body rows): keep entire table together
    # by setting keepWithNext on all cells' paragraphs except last row
    if total_rows <= 6:  # 1 header + ≤5 body rows
        for idx, row in enumerate(t.rows):
            if idx < total_rows - 1:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.keep_with_next = True

    _spacer(doc, 4)  # compact spacing after table
    return t


def _section_title(doc, num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{num}  {title}")
    _set_run(run, size=16, bold=True, color=BRAND)
    _add_horizontal_line(p, color=ACCENT)


def _subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_run(run, size=11, bold=True, color=BRAND_LIGHT)


def _para(doc, text, *, size=10, color=None, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run(run, size=size, color=color)
    if italic:
        run.italic = True
    return p


def _bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    if p.runs:
        p.runs[0].text = text
        run = p.runs[0]
    else:
        run = p.add_run(text)
    _set_run(run, size=9.5)


def _highlight_box(doc, text, bg_color=HIGHLIGHT_BG):
    """Add a highlighted paragraph with background color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    _set_run(run, size=10, color=BRAND)
    # Set paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg_color)
    pPr.append(shd)


def _spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pts)
    p.paragraph_format.space_after = Pt(0)
    pf = p.paragraph_format
    pf.line_spacing = Pt(pts)


def _s(v, prefix="", suffix="", default="—"):
    """Safe string: return default if v is None/empty, else prefix+str(v)+suffix."""
    if v is None or str(v).strip() in ("", "None"):
        return default
    return f"{prefix}{v}{suffix}"


# ── Main report builder ──
def build_report(state) -> str:
    doc = Document()

    # Page margins & footer
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

        # Page number in footer
        footer = s.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        run = fp.add_run()
        _set_run(run, size=8, color=MUTED)
        # Insert PAGE field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        run._r.append(fldChar2)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar3)
        run2 = fp.add_run(" / ")
        _set_run(run2, size=8, color=MUTED)
        run3 = fp.add_run()
        _set_run(run3, size=8, color=MUTED)
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "begin")
        run3._r.append(fldChar4)
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(qn("xml:space"), "preserve")
        instrText2.text = " NUMPAGES "
        run3._r.append(instrText2)
        fldChar5 = OxmlElement("w:fldChar")
        fldChar5.set(qn("w:fldCharType"), "separate")
        run3._r.append(fldChar5)
        fldChar6 = OxmlElement("w:fldChar")
        fldChar6.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar6)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    keywords = state.get("keywords", [])
    asins = state.get("asins", [])
    market = state.get("market", {})
    market_overview = state.get("market_overview", {})
    market_trends = state.get("market_trends", [])
    comp = state.get("competition", {})
    pricing = state.get("pricing", {})
    bad = state.get("bad_reviews", {})
    directions = state.get("directions", {}).get("items", [])
    conclusion = state.get("conclusion", [])
    warnings = state.get("warnings", [])
    data_stats = state.get("data_stats", {})
    keyword_analysis = state.get("keyword_analysis", [])
    category_trends = state.get("category_trends", [])
    source_conclusions = state.get("source_conclusions", [])
    insights = state.get("insights", {})
    profit_calc = state.get("profit_calc", [])
    total_products = data_stats.get("total_products", 0)
    total_reviews = data_stats.get("total_reviews", 0)

    # ═════════════════════════════════════════
    # COVER
    # ═════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("Amazon 选品评估报告")
    _set_run(tr, size=30, bold=True, color=BRAND)

    _spacer(doc, 4)

    if keywords:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub_p.add_run(" · ".join(keywords[:3]))
        _set_run(sr, size=14, color=ACCENT, bold=True)

    bar = doc.add_paragraph()
    bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br_ = bar.add_run("━" * 30)
    _set_run(br_, size=10, color=ACCENT)

    # Meta info
    meta_lines = [
        f"数据截至  {datetime.now().strftime('%Y-%m-%d')}",
        f"分析 ASIN  {len(asins)} 个",
        f"目标利润率  {int(state.get('target_margin', 0.3)*100)}%",
    ]
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta_p.add_run("    |    ".join(meta_lines))
    _set_run(mr, size=9.5, color=MUTED)

    src_p = doc.add_paragraph()
    src_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src_text = f"数据来源: 上传 Excel（{total_products} 个产品"
    if total_reviews:
        src_text += f"，{total_reviews} 条评论"
    src_text += "）"
    src_run = src_p.add_run(src_text)
    _set_run(src_run, size=9, color=MUTED)
    src_run.italic = True

    if warnings:
        _spacer(doc, 4)
        for w in warnings:
            wp = doc.add_paragraph()
            wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            wr = wp.add_run(f"⚠ {w}")
            _set_run(wr, size=9, color=DANGER)

    doc.add_page_break()

    # ═════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ═════════════════════════════════════════
    exec_summary = insights.get("executive_summary", "")
    if exec_summary:
        _section_title(doc, "◆", "执行摘要")
        _highlight_box(doc, exec_summary)
        _spacer(doc, 8)

    # ═════════════════════════════════════════
    # 一、MARKET OVERVIEW (expanded)
    # ═════════════════════════════════════════
    _section_title(doc, "一", "类目市场概览")
    _para(doc,
        f"基于上传 Excel 中 {total_products} 个产品数据，"
        "从市场容量、价格分布、竞争格局等维度评估市场。",
        size=9.5, color=MUTED,
    )

    # ── 1.1 Market basics table ──
    _subtitle(doc, "▎ 1.1 市场基本面")
    market_keys = list(market.keys())
    if market_keys:
        headers_m = ["指标"] + [k[:20] for k in market_keys]
        rows_m = [
            ["匹配产品数"] + [_s(market.get(k, {}).get("result_count")) for k in market_keys],
            ["价格区间"] + [
                f"{_s(market.get(k,{}).get('price_min'), '$')} – {_s(market.get(k,{}).get('price_max'), '$')}"
                for k in market_keys
            ],
            ["中位价"] + [_s(market.get(k,{}).get('price_median'), '$') for k in market_keys],
            ["销量带"] + [_s(market.get(k, {}).get("sales_band")) for k in market_keys],
            ["热度评级"] + [_s(market.get(k, {}).get("rating")) for k in market_keys],
            ["趋势"] + [_s(market.get(k, {}).get("trend")) for k in market_keys],
        ]
        _add_table(doc, headers_m, rows_m)
    else:
        _para(doc, "未输入关键词，无关键词维度市场分析。", color=MUTED, italic=True)

    # ── 1.2 Market concentration & overview (only show if enough data) ──
    if market_overview:
        overview_rows = []
        _ov = market_overview
        if _ov.get("monthly_total_sales"):
            overview_rows.append(["月总销量", f"{_ov['monthly_total_sales']:,}"])
        if _ov.get("monthly_total_revenue"):
            overview_rows.append(["月总销售额", f"${_ov['monthly_total_revenue']:,.0f}"])
        if _ov.get("avg_price"):
            overview_rows.append(["平均价格", f"${_ov['avg_price']:.2f}"])
        if _ov.get("avg_rating"):
            overview_rows.append(["平均星级", f"{_ov['avg_rating']}"])
        if _ov.get("avg_reviews"):
            overview_rows.append(["平均评论数", f"{_ov['avg_reviews']:,}"])
        if _ov.get("sample_count"):
            text = f"{_ov['sample_count']} 个商品"
            if _ov.get("brand_count"):
                text += f"，{_ov['brand_count']} 个品牌"
            if _ov.get("seller_count"):
                text += f"，{_ov['seller_count']} 个卖家"
            overview_rows.append(["样本规模", text])
        if _ov.get("brand_concentration"):
            overview_rows.append(["品牌集中度", _ov["brand_concentration"]])
        if _ov.get("seller_concentration"):
            overview_rows.append(["卖家集中度", _ov["seller_concentration"]])
        if _ov.get("product_concentration"):
            overview_rows.append(["商品集中度", _ov["product_concentration"]])
        if _ov.get("new_product_pct"):
            overview_rows.append(["新品占比", _ov["new_product_pct"]])
        if _ov.get("cn_seller_pct"):
            overview_rows.append(["中国卖家占比", _ov["cn_seller_pct"]])
        if _ov.get("fba_pct"):
            overview_rows.append(["FBA 卖家占比", _ov["fba_pct"]])
        if _ov.get("top_brands_text"):
            overview_rows.append(["头部品牌", _ov["top_brands_text"][:60]])
        if len(overview_rows) >= 3:
            _subtitle(doc, "▎ 1.2 市场集中度与竞争格局")
            _add_table(doc, ["指标", "数值"], overview_rows)

    # ── 1.3 Monthly trends ──
    if market_trends:
        _subtitle(doc, "▎ 1.3 月度趋势")
        # Pick key columns for trend display
        trend_cols = ["月份"]
        possible = ["月总销量", "月总销售额($)", "月总销售额", "月均销量",
                     "平均价格($)", "平均价格", "平均星级", "新品占比",
                     "品牌集中度", "卖家集中度"]
        for col in possible:
            if any(col in t for t in market_trends):
                trend_cols.append(col)
        if len(trend_cols) > 1:
            trend_rows = []
            for t in market_trends[:12]:  # max 12 months
                row = [_s(t.get("月份"))]
                for col in trend_cols[1:]:
                    v = t.get(col)
                    if v is not None:
                        try:
                            fv = float(str(v).replace(",", "").replace("$", ""))
                            if fv < 1 and col in ("新品占比", "品牌集中度", "卖家集中度"):
                                row.append(f"{fv*100:.1f}%")
                            elif fv > 10000:
                                row.append(f"{fv:,.0f}")
                            else:
                                row.append(f"{fv:,.2f}" if fv != int(fv) else str(int(fv)))
                        except (ValueError, TypeError):
                            row.append(str(v)[:20])
                    else:
                        row.append("—")
                trend_rows.append(row)
            if trend_rows:
                _add_table(doc, trend_cols, trend_rows)

    # ── 1.4 Category trends (text) ──
    if category_trends:
        _subtitle(doc, "▎ 1.4 类目趋势")
        for ct in category_trends[:5]:
            text = str(ct).strip()
            if len(text) > 200:
                _para(doc, text, size=9.5)
            elif text:
                _bullet(doc, text)

    market_insights = insights.get("market_insights", [])
    if market_insights:
        _subtitle(doc, "▎ 市场洞察")
        for mi in market_insights:
            _bullet(doc, str(mi))

    # ═════════════════════════════════════════
    # 二、KEYWORD ANALYSIS (new section)
    # ═════════════════════════════════════════
    if keyword_analysis:
        doc.add_page_break()
        _section_title(doc, "二", "关键词分析")
        _para(doc,
            "核心关键词的搜索量、转化率、CPC 竞价与竞争程度，为广告投放提供数据支撑。",
            size=9.5, color=MUTED,
        )

        _subtitle(doc, "▎ 2.1 核心关键词总览")
        # Build keyword table from keyword_analysis data
        kw_headers = ["关键词"]
        # Detect available columns
        sample = keyword_analysis[0] if keyword_analysis else {}
        kw_col_map = [
            ("翻译", "翻译"),
            ("月搜索量", "月搜索量"), ("月购买量", "月购买量"),
            ("转化率", "转化率"), ("点击集中度", "点击集中度"),
            ("CPC精确($)", "CPC精确($)"), ("CPC泛($)", "CPC泛($)"),
            ("竞争程度", "竞争程度"), ("搜索趋势", "搜索趋势"),
        ]
        active_cols = []
        for label, key in kw_col_map:
            if key in sample or any(key in kw for kw in keyword_analysis[:5]):
                kw_headers.append(label)
                active_cols.append(key)
        kw_rows = []
        for kw in keyword_analysis[:20]:
            kw_name = kw.get("关键词") or kw.get("keyword") or kw.get("序号", "")
            row = [str(kw_name)[:30]]
            for key in active_cols:
                row.append(_s(kw.get(key)))
            kw_rows.append(row)
        if kw_rows:
            _add_table(doc, kw_headers, kw_rows)

        # Keyword strategy insights from LLM
        kw_strategy = insights.get("keyword_strategy", [])
        if kw_strategy:
            _subtitle(doc, "▎ 2.2 关键词投放策略")
            for ks in kw_strategy:
                _bullet(doc, str(ks))

    # ═════════════════════════════════════════
    # 三、COMPETITION (expanded, merged with selling points comparison)
    # ═════════════════════════════════════════
    doc.add_page_break()
    _section_title(doc, "三", "竞品深度分析")

    comp_rows = comp.get("rows", [])
    if comp_rows:
        _subtitle(doc, "▎ 3.1 竞品矩阵")
        rows2 = []
        for r in comp_rows:
            rows2.append([
                r.get("asin"), _s(r.get("brand")),
                _s(r.get("price"), "$"), _s(r.get("rating")),
                _s(r.get("reviews_total")), _s(r.get("est_monthly_sales")),
                _s(r.get("competitiveness")),
            ])
        _add_table(doc, ["ASIN", "品牌", "价格", "评分", "评论数", "月销量", "竞争力"], rows2)

    heat = comp.get("heat", {})
    if heat:
        _subtitle(doc, "▎ 3.2 关键词竞争热度")
        rows_h = [[k, h.get("热度", "—"), h.get("价格区间", "—"), h.get("竞争激烈度", "—")] for k, h in heat.items()]
        _add_table(doc, ["关键词", "热度", "价格区间", "竞争激烈度"], rows_h)

    # ── 3.3 Selling points vs pain points comparison (merged from old section 五) ──
    comp_comparison = bad.get("comp_comparison", [])
    if comp_comparison:
        _subtitle(doc, "▎ 3.3 竞品卖点 vs 差评对比")
        _para(doc,
            "基于竞品 Listing 卖点与消费者实际差评的对比，找出宣传与现实的落差。",
            size=9.5, color=MUTED,
        )
        rows_cc = []
        for cc in comp_comparison:
            product_type = cc.get("product_type") or cc.get("brand", "")
            selling = "\n".join(f"• {s}" for s in cc.get("selling_points", [])[:3]) if cc.get("selling_points") else cc.get("positive_keywords", "—")
            pain = cc.get("negative_keywords") or "—"
            weakness = cc.get("weakness") or ""
            if weakness:
                pain += f"\n• {weakness}"
            rows_cc.append([product_type, selling, pain])
        _add_table(doc, ["产品类型", "宣传卖点（正面）", "实际差评痛点（负面）"], rows_cc)

        improvements = [cc for cc in comp_comparison if cc.get("improvement")]
        if improvements:
            _subtitle(doc, "▎ 关键改进机会")
            for cc in improvements:
                _bullet(doc, f"{cc.get('brand', '')}({cc.get('asin', '')}): {cc['improvement']}")

    comp_insights = insights.get("competition_insights", [])
    if comp_insights:
        _subtitle(doc, "▎ 竞争洞察")
        for ci in comp_insights:
            _bullet(doc, str(ci))

    # ═════════════════════════════════════════
    # 四、BAD REVIEWS
    # ═════════════════════════════════════════
    doc.add_page_break()
    total_critical = bad.get("total_critical", 0)
    if total_critical > 0:
        _section_title(doc, "四", "竞品差评深度分析")
        _para(doc,
            f"共分析 {total_critical} 条评论（去重后 {bad.get('unique_used', 0)} 条），"
            f"基于 {total_reviews} 条上传评论数据。",
            size=9.5, color=MUTED,
        )
    else:
        _section_title(doc, "四", "市场痛点推断")
        _para(doc,
            "未上传评论数据，以下基于产品特征推断，仅供参考。",
            color=MUTED, italic=True,
        )

    overall_summary = bad.get("overall_summary", "")
    if overall_summary:
        _subtitle(doc, "▎ 4.1 差评整体分析")
        _highlight_box(doc, overall_summary)

    top10 = bad.get("top10", [])
    if top10:
        _subtitle(doc, "▎ 4.2 TOP 痛点明细")
        rows4 = []
        for it in top10:
            rows4.append([
                it.get("rank", ""),
                it.get("issue", ""),
                it.get("frequency", ""),
                it.get("dimension", ""),
                it.get("severity", "—"),
                (it.get("typical_quote") or "")[:80],
            ])
        _add_table(doc, ["#", "痛点", "频次", "维度", "严重度", "典型原句"], rows4)

        # Root cause + suggestion details
        _subtitle(doc, "▎ 4.3 根因分析与改良建议")
        for it in top10:
            rc = it.get("root_cause", "")
            sg = it.get("suggestion", "")
            fs = it.get("feasibility", "")
            if rc or sg:
                issue = it.get("issue", "")
                rank = it.get("rank", "")
                detail = f"{rank}. {issue}"
                if rc:
                    detail += f"\n   根因：{rc}"
                if sg:
                    detail += f"\n   建议：{sg}"
                if fs:
                    detail += f"（可行性：{fs}）"
                _bullet(doc, detail)

    # Review excerpts
    by_kw = bad.get("by_keyword", {})
    if any(by_kw.values()):
        _subtitle(doc, "▎ 4.4 差评原句摘录")
        for kw, bullets in by_kw.items():
            if bullets:
                p = doc.add_paragraph()
                p.paragraph_format.keep_with_next = True
                r = p.add_run(f"  {kw}")
                _set_run(r, size=10, bold=True, color=BRAND_LIGHT)
                for b in bullets:
                    _bullet(doc, b)

    # ═════════════════════════════════════════
    # 五、PRICING & PROFIT (expanded)
    # ═════════════════════════════════════════
    doc.add_page_break()
    _section_title(doc, "五", "利润测算与定价策略")
    pricing_keys = list(pricing.keys())
    has_excel_profit = any(pricing.get(k, {}).get("source") == "Excel 利润核算" for k in pricing_keys)

    # ── 5.1 Full cost breakdown from Excel profit_calc ──
    if profit_calc:
        _subtitle(doc, "▎ 5.1 成本结构拆解")
        _para(doc, "数据来自上传 Excel 利润核算表，展示完整成本结构。", size=9.5, color=MUTED)
        # profit_calc is list of dicts; display as table
        if isinstance(profit_calc[0], dict):
            # Get all unique keys for headers
            all_keys = []
            for row in profit_calc[:10]:
                for k in row.keys():
                    if k not in all_keys:
                        all_keys.append(k)
            # Filter out empty columns
            active_keys = []
            for k in all_keys:
                if any(row.get(k) is not None and str(row.get(k, "")).strip() not in ("", "None") for row in profit_calc[:10]):
                    active_keys.append(k)
            if active_keys:
                pc_headers = [str(k)[:15] for k in active_keys]
                pc_rows = []
                for row in profit_calc[:15]:
                    pc_rows.append([_s(row.get(k)) for k in active_keys])
                _add_table(doc, pc_headers, pc_rows)

    # ── 5.2 Pricing recommendation ──
    _subtitle(doc, "▎ 5.2 定价建议")
    if has_excel_profit:
        _para(doc, "基于 Excel 利润核算数据的定价方案。", size=9.5, color=MUTED)
        rows3 = []
        for k in pricing_keys:
            p = pricing.get(k, {})
            rows3.append([
                k[:20],
                _s(p.get("sell_price_usd"), "$"),
                _s(p.get("product_cost_cny"), "¥"),
                _s(p.get("shipping_cost_cny"), "¥"),
                _s(p.get("fba_total_usd"), "$"),
                _s(p.get("margin_rate")),
                _s(p.get("gross_profit_cny"), "¥"),
            ])
        _add_table(doc,
            ["款式/方式", "售价", "产品成本", "头程运费", "FBA成本", "毛利率", "毛利润"],
            rows3,
        )
    elif pricing_keys:
        _para(doc,
            f"基于市场中位价估算；目标毛利率 {int(state.get('target_margin',0.3)*100)}%，"
            f"FBA 费率 {int(state.get('fee_rate',0.3)*100)}%。",
            size=9.5, color=MUTED,
        )
        rows3 = []
        for k in pricing_keys:
            p = pricing.get(k, {})
            rows3.append([
                k, _s(p.get("median_price"), "$"), _s(p.get("entry_price"), "$"),
                _s(p.get("entry_range")), _s(p.get("target_cost_max"), "$"),
                _s(p.get("expected_gross_per_unit"), "$"),
            ])
        _add_table(doc,
            ["关键词", "中位价", "入场价", "价格区间", "成本上限", "单件毛利"],
            rows3,
        )

    pricing_insights = insights.get("pricing_insights", [])
    if pricing_insights:
        _subtitle(doc, "▎ 5.3 定价策略洞察")
        for pi in pricing_insights:
            _bullet(doc, str(pi))

    # ═════════════════════════════════════════
    # 六、DIRECTIONS
    # ═════════════════════════════════════════
    doc.add_page_break()
    _section_title(doc, "六", "产品上新方向")

    for di, d in enumerate(directions):
        # ── Separator between directions ──
        if di > 0:
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(6)
            sep.paragraph_format.space_after = Pt(6)
            _add_horizontal_line(sep, color=BORDER, sz="4")

        # ── Direction title: "方向一：xxx" ──
        label = {0: "方向一", 1: "方向二", 2: "方向三"}.get(di, f"方向{di+1}")
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(14)
        title_p.paragraph_format.space_after = Pt(4)
        title_p.paragraph_format.keep_with_next = True
        r1 = title_p.add_run(f"{label}：")
        _set_run(r1, size=13, bold=True, color=BRAND)
        r2 = title_p.add_run(d.get("name", "—"))
        _set_run(r2, size=13, bold=True, color=BRAND_LIGHT)

        # ── Positioning ──
        positioning = d.get("positioning", "—")
        if positioning and positioning != "—":
            _para(doc, positioning, size=10)

        # ── Key metrics: single compact line ──
        metrics = []
        if d.get("target_user"):
            metrics.append(f"目标人群：{d['target_user']}")
        if d.get("target_price"):
            metrics.append(f"目标价：${d['target_price']}")
        if d.get("monthly_sales_target"):
            metrics.append(f"月销目标：{d['monthly_sales_target']}")
        if metrics:
            mp = doc.add_paragraph()
            mp.paragraph_format.space_before = Pt(2)
            mp.paragraph_format.space_after = Pt(6)
            mr = mp.add_run("    |    ".join(metrics))
            _set_run(mr, size=9.5, color=MUTED)

        # ── Market opportunity ──
        if d.get("market_opportunity"):
            _subtitle(doc, "▎ 市场机会")
            _para(doc, d["market_opportunity"], size=10)

        # ── Evidence ──
        if d.get("evidence"):
            _subtitle(doc, "▎ 数据依据")
            for ev in d["evidence"]:
                _bullet(doc, str(ev))

        # ── Improvements as 2-column table [痛点改进 | 差异化功能] ──
        improvements_list = d.get("improvements") or []
        if improvements_list:
            _subtitle(doc, "▎ 改良点")
            mid = (len(improvements_list) + 1) // 2
            col1 = improvements_list[:mid]
            col2 = improvements_list[mid:]
            left = "\n".join(f"• {s}" for s in col1)
            right = "\n".join(f"• {s}" for s in col2) if col2 else "—"
            _add_table(doc, ["针对差评痛点的改进", "差异化新增功能"], [[left, right]])

        # ── Cost estimate ──
        if d.get("cost_estimate"):
            _subtitle(doc, "▎ 成本与利润预估")
            _para(doc, d["cost_estimate"], size=10)

        # ── Risks ──
        if d.get("risks"):
            _subtitle(doc, "▎ 风险提示")
            for rk in d["risks"]:
                _bullet(doc, f"⚠ {rk}")

    # ═════════════════════════════════════════
    # 七、CONCLUSION
    # ═════════════════════════════════════════
    doc.add_page_break()
    _section_title(doc, "七", "综合结论与行动优先级")

    # Category summary from insights
    cat_summary = insights.get("category_summary", "")
    if cat_summary:
        _subtitle(doc, "▎ 类目综合评估")
        _highlight_box(doc, cat_summary)

    if conclusion:
        _subtitle(doc, "▎ 行动优先级")
        rows7 = [[c.get("medal", ""), c.get("direction", ""), c.get("reason", ""),
                   c.get("target_price", ""), c.get("next_step", "")] for c in conclusion]
        _add_table(doc, ["优先级", "方向", "切入理由", "目标定价", "建议下一步"], rows7)

    # Source conclusions reference
    if source_conclusions:
        _subtitle(doc, "▎ Excel 原始结论（参考）")
        for sc in source_conclusions[:5]:
            text = str(sc).strip()
            if text:
                _bullet(doc, text)

    # ═════════════════════════════════════════
    # FOOTER
    # ═════════════════════════════════════════
    _spacer(doc, 20)
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = line.add_run("━" * 30)
    _set_run(lr, size=8, color=ACCENT)

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(f"— 报告结束  |  生成时间 {datetime.now().strftime('%Y-%m-%d %H:%M')} —")
    _set_run(fr, size=9, color=MUTED)
    fr.italic = True

    main_kw = (keywords[0] if keywords else "report").replace(" ", "_")
    fname = f"Amazon选品评估报告_{main_kw}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    fpath = os.path.join(OUTPUT_DIR, fname)
    doc.save(fpath)
    return fpath
